import os
import magic
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
import pythoncom

def convert_file(input_path, target_format):
    try:
        # Determine input file type
        file_type = magic.from_file(input_path, mime=True)
        
        filename = os.path.basename(input_path)
        name_without_ext = os.path.splitext(filename)[0]
        output_filename = f"{name_without_ext}_converted.{target_format}"
        output_path = os.path.join(os.path.dirname(input_path), output_filename)
        
        # PDF conversions
        if file_type == 'application/pdf':
            if target_format in ['jpg', 'jpeg', 'png']:
                images = convert_from_path(input_path)
                if images:
                    images[0].save(output_path, format=target_format.upper())
                    return output_path, output_filename
                    
            elif target_format == 'txt':
                text = pytesseract.image_to_string(Image.open(input_path))
                with open(output_path, 'w') as f:
                    f.write(text)
                return output_path, output_filename
                
        # Image conversions
        elif file_type.startswith('image/'):
            img = Image.open(input_path)
            
            if target_format in ['jpg', 'jpeg', 'png']:
                img.save(output_path, format=target_format.upper())
                return output_path, output_filename
                
            elif target_format == 'pdf':
                img.save(output_path, "PDF", resolution=100.0)
                return output_path, output_filename
                
            elif target_format == 'txt':
                text = pytesseract.image_to_string(img)
                with open(output_path, 'w') as f:
                    f.write(text)
                return output_path, output_filename
                
        # Word document conversions
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword']:
            pythoncom.CoInitialize()
            if target_format == 'pdf':
                # This requires Microsoft Word installed or a cloud service
                # For simplicity, we'll just convert to text here
                doc = Document(input_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                output_path = output_path.replace('.pdf', '.txt')
                with open(output_path, 'w') as f:
                    f.write(text)
                return output_path, output_filename.replace('.pdf', '.txt')
                
            elif target_format == 'txt':
                doc = Document(input_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                with open(output_path, 'w') as f:
                    f.write(text)
                return output_path, output_filename
                
        # Text file conversions
        elif file_type == 'text/plain':
            if target_format == 'docx':
                doc = Document()
                with open(input_path, 'r') as f:
                    doc.add_paragraph(f.read())
                doc.save(output_path)
                return output_path, output_filename
        
        return None, None
        
    except Exception as e:
        print(f"Conversion error: {e}")
        return None, None